"""
This script defines an algorithm to generate English-Hindi sentence pairs 
suitable for training a translation algorithm.

This approach leverages a provided translation function and a similarity metric 
to identify the closest Hindi sentence to the translated version of each English sentence.

**Inputs:**

* `en_file`: Path to the source language (English) file.
* `hi_file`: Path to the target language (Hindi) file.
* `translate_fn`: A function that translates text from source to target language (e.g., English to Hindi).

**Output:**

* A list of English-Hindi sentence pairs suitable for training a translation algorithm.

**Algorithm:**

1. **Load Text Files:**
   * Read the content of `en_file` and store each line (sentence) in a list called `en_sentences`.
   * Read the content of `hi_file` and store each line (sentence) in a list called `hi_sentences`.

2. **Translate English Sentences:**
   * Create a new list called `target_en_sentences`.
   * For each sentence `en_sentence` in `en_sentences`:
       * Use `translate_fn` to translate `en_sentence` to Hindi.
       * Add the translated sentence to `target_en_sentences`.

3. **Iterate through English Sentences:**
   * For each sentence `en_sentence` (and its corresponding translation `target_en_sentence`) in `en_sentences` and `target_en_sentences`:

       **Find Closest Hindi Match:**
           * Initialize `closest_hi_score` to negative infinity (very low score).
           * Initialize `closest_hi_sentence` to `None`.
           * For each sentence `hi_sentence` in `hi_sentences`:
               * Calculate the similarity score between `target_en_sentence` and `hi_sentence` using the chosen similarity function (explained later). Let's call this score `score`.
               * If `score` is greater than `closest_hi_score`:
                   * Update `closest_hi_score` with `score`.
                   * Update `closest_hi_sentence` with `hi_sentence`.

       **Add Pair to Output:**
           * Add a tuple containing `en_sentence` and `closest_hi_sentence` to the output list.

**Similarity Function:**

This algorithm relies on a similarity function to compare translated English sentences (`target_en_sentences`) and Hindi sentences (`hi_sentences`). Here are two options:

* **Simple Word Overlap:** Calculate the number of words that appear in both sentences.
* **More Advanced Techniques:** Use libraries or tools that provide semantic similarity scores between sentences, considering word meaning and context.

**Note:**

This algorithm finds the closest Hindi sentence based on the similarity between the translated English sentence and the original Hindi sentence.
Adjustments might be needed to handle sentences with different lengths.

**Output List:**

The final output will be a list of tuples, where each tuple contains an English sentence and its corresponding closest Hindi sentence based on the chosen similarity metric. This list can be used to train a translation algorithm.
"""

import docx

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import sys
from tokenizer import get_sentences
import json
from translation import translate_text

def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def get_sentences_from_docx(filename, language) -> list:
    """
    Extracts sentences from a docx file and returns a list of sentences.
    """
    document = docx.Document(filename)
    source_paras = []
    source_sentences = []

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            source_paras.append(block.text)

        elif isinstance(block, Table):
            for i,row in enumerate(block.rows):
                for j,cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        # source_paras.append(cell.paragraphs[0])
                        for run in paragraph.runs:
                            source_paras.append(run.text)

    # source_strings for each source_string use get_sentences_en and append that to the array
    for source_string in source_paras:
        sentences = get_sentences(source_string, language)
        source_sentences.extend(sentences)
    source_sentences = list(map(lambda x: x.strip().replace("\xa0","").replace("\t", "  "), source_sentences))
    source_sentences = list(filter(lambda x: len(x)>15, source_sentences))
    return source_sentences

def process_doc(source_filename, target_filename, output_filename, source_language, target_language):
    source_sentences=get_sentences_from_docx(source_filename, source_language)
    # translated_sentences = list of map of each value of source_sentences using translate_text function
    translated_sentences = [translate_text(sentence) for sentence in source_sentences]
    target_sentences=get_sentences_from_docx(target_filename, target_language)

    # dump source_sentences, translated_sentences, target_sentences in a json file
    data = {
        "source_sentences": source_sentences,
        "translated_sentences": translated_sentences,
        "target_sentences": target_sentences
    }

    # add encoding to open file syntax below

    with open('sentences.json', 'w', encoding='utf-8') as f:
        json.dump(data, f,  ensure_ascii=False)

    # read the json file and print source_sentences, translated_sentences, target_sentences
    with open('sentences.json', 'r', encoding="utf-8") as f:
        data = json.load(f)
        source_sentences = data["source_sentences"]
        translated_sentences = data["translated_sentences"]
        target_sentences = data["target_sentences"]
        
    # with open(output_filename, 'w', encoding='utf-8') as f:
    #     for source_sentence in source_sentences:
    #         translated_sentences.append(translate_paragraph("google", source_sentence, target_language))
    #     for source_sentence in source_sentences:
    #         closest_hi_score = -1
    #         closest_hi_sentence = None
    #         for target_sentence in translated_sentences:
    #             for hi_sentence in target_sentences:
    #                 score = similarity(target_sentence, hi_sentence)
    #                 if score > closest_hi_score:
    #                     closest_hi_score = score
    #                     closest_hi_sentence = hi_sentence
    #         f.write(source_sentence + "\t" + closest_hi_sentence + "\n")
    # print(len(source_sentences), len(target_sentences))

if __name__ == '__main__':
    # Check if the correct number of command line arguments is provided
    if len(sys.argv) < 3:
        print("Usage: python script.py source_file target_file [output_file] [source_language] [target_language]")
        sys.exit(1)

    # Extract the input and output file names from command line arguments
    source_file = sys.argv[1]
    target_file = sys.argv[2]

    # Extract the target language if provided,default to 'en' if not specified
    output_file = sys.argv[3] if len(sys.argv) > 3 else 'translation_pairs.tsv'
    source_language = sys.argv[4] if len(sys.argv) > 4 else 'en'
    target_language = sys.argv[5] if len(sys.argv) > 3 else 'hi'

    # Call the function to output the translation
    process_doc(source_file, target_file, output_file, source_language, target_language)