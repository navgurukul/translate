# Create a script in python to read a docx file with one paragraph after other. paragraphs are alternatively in hindi and english. create a tsv file from this, where the english paragraph is in the first column, and the second column is hindi paragraph. use a simple language detection tool which can be based on the characters used in the text to identify which is hindi and which is english.

import docx
import os
import sys
import re
import langid
import argparse
import nltk.data

tokenizer = nltk.data.load('tokenizers/punkt/english.pickle')

# A function that inputs a paragraph and returns an array of the sentences in the paragraph
def get_sentences_en(paragraph):
    # Split the paragraph into sentences
    sentences = tokenizer.tokenize(paragraph)
    # re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s', paragraph)
    return sentences

# A function that inputs a paragraph in Hindi and returns an array of the sentences in the paragraph
def get_sentences_hi(paragraph):
    # Split the paragraph into sentences
    sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=।|\?)\s', paragraph)
    return sentences

def clean_paragraphs(paragraphs):
    # map the above paragraphs text to a new variable paragraphs_text
    paragraphs = [paragraph.text.strip() for paragraph in paragraphs]
    paragraphs = [paragraph.replace('\t','') for paragraph in paragraphs]
    
    # map the paragraphs to remove the first character of the paragraphs beginning with — or -
    paragraphs = [re.sub(r'^—', '', paragraph) for paragraph in paragraphs]
    paragraphs = [re.sub(r'^-', '', paragraph) for paragraph in paragraphs]

    # filter out the empty paragraphs
    paragraphs = list(filter(lambda x: x != '', paragraphs))
    paragraphs = list(filter(lambda x: x != '—', paragraphs))

    return paragraphs

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input_file_en", help="input file english")
    parser.add_argument("input_file_hi", help="input file hindi")
    parser.add_argument("output_file", help="output file")
    args = parser.parse_args()

    input_file_en = args.input_file_en
    input_file_hi = args.input_file_hi
    output_file = args.output_file

    doc_en = docx.Document(input_file_en)
    doc_hi = docx.Document(input_file_hi)

    paragraphs_en = clean_paragraphs(doc_en.paragraphs)
    paragraphs_hi = clean_paragraphs(doc_hi.paragraphs)

    g = open("unmatched.tsv", 'w', encoding='utf-8')

    with open(output_file, 'w', encoding='utf-8') as f:  # Specify encoding as utf-8

        # Initialize a variable i to 0
        i = 0

        # Add an assertion to verify that the length of paragraphs_en is equal to the length of paragraphs_hi and print the length of the paragraphs if the assertion fails
        assert len(paragraphs_en) == len(paragraphs_hi), "Length of paragraphs_en is not equal to length of paragraphs_hi. Length of paragraphs_en is " + str(len(paragraphs_en)) + " and length of paragraphs_hi is " + str(len(paragraphs_hi))
    
        # Loop while i is less than max_len
        while i < len(paragraphs_en):
            # Do the same operations as in the for loop
            en_sentences = get_sentences_en(paragraphs_en[i])
            hi_sentences = get_sentences_hi(paragraphs_hi[i])

            if len(en_sentences) == len(hi_sentences):
                for j in range(len(en_sentences)):
                    f.write(en_sentences[j] + '\t' + hi_sentences[j]+'\n')
            else:
                # write these cases in a unmatched.tsv file
                for j in range(max(len(en_sentences),len(hi_sentences))):
                    if j < min(len(en_sentences), len(hi_sentences)):
                        g.write(en_sentences[j] + '\t' + hi_sentences[j]+'\n')
                    if j >= len(en_sentences):
                        g.write('\t' + hi_sentences[j]+'\n')
                    if j >= len(hi_sentences):
                        g.write(en_sentences[j] + '\t'+'\n')

                g.write('\n')
                # print(en_sentences, hi_sentences,'\n')
            
            i += 1
    
    g.close()

if __name__ == '__main__': 
    main()