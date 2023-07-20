import csv
import sys
from docx import Document
from docx.shared import RGBColor

post_replacements_file = 'post-phrases.csv'

with open(post_replacements_file, 'r', encoding='utf-8') as csvfile:
    reader = csv.reader(csvfile)
    post_replacements = {row[0]: row[1] for row in reader}

def replace_phrases(text):
    for phrase, replacement in post_replacements.items():
        text = text.replace(phrase, replacement)

    return text

# Get color of a paragraph from a docx file
def get_paragraph_color(paragraph):
    if paragraph.runs:
        first_run = paragraph.runs[0]
        if hasattr(first_run, 'font'):
            font_color = first_run.font.color
            if font_color is not None and isinstance(font_color.rgb, tuple):
                return RGBColor(*font_color.rgb)
    return None

def post_replace_phrases(input_path, output_path):
    doc = Document(input_path)
    output_doc = Document()

    for paragraph in doc.paragraphs:
        text=paragraph.text.strip()
        if text=='':
            continue

        updated_paragraph = replace_phrases(paragraph.text).replace('\n', '')
        run = output_doc.add_paragraph().add_run()
        run.text = updated_paragraph
        font = run.font
        font.color.rgb = get_paragraph_color(paragraph)

    output_doc.save(output_path)

if __name__ == '__main__':
    # Check if the correct number of command line arguments is provided
    if len(sys.argv) < 2:
        print("Usage: python post-phrases-edit.py input_file output_file")
        sys.exit(1)

    # Extract the input and output file names from command line arguments
    input_file = sys.argv[1]
    output_file = sys.argv[2]

    # Call the function to output the translation
    post_replace_phrases(input_file, output_file)
