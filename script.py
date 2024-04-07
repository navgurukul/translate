#!/usr/bin/env python
# encoding: utf-8

'''
This Python script is designed to read a Microsoft Word document (.docx), translate its content to a specified language, and save the translated content into a new Word document. The script uses the python-docx library to interact with Word documents and the tqdm library to display a progress bar.

The process_doc(input_filename, output_filename, target_language) function reads a Word document, translates its content, and writes the translated content to a new document. It first opens the input document and creates a new document. It then creates a progress bar with tqdm, setting the total to the number of paragraphs in the document.

The function then iterates over each block (paragraph or table) in the document. If the block is a paragraph, it adds a new paragraph to the new document and copies the text and formatting from the original paragraph. It then adds a translated version of the paragraph text to the new paragraph, setting the font color to green.

If the block is a table, it adds a new table to the new document with twice as many rows as the original table (to accommodate the translated text) and the same number of columns. It then creates a 2D list to store the text and formatting of each cell in the original table. It iterates over each cell in the original table, concatenating the text of all runs in the cell and storing the text and formatting in the 2D list. It then iterates over the cells in the new table, setting the text of each cell to the corresponding text in the 2D list (translated if the row is odd).

Finally, the function updates the progress bar, closes it when done, and saves the new document.
'''

from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from translation import translate_text

import sys
from docx import Document
import random
from docx.shared import RGBColor
from tqdm import tqdm

def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def process_doc(input_filename, output_filename, target_language):
    """
    Reads a docx file in order, replacing characters with spaces randomly,
    and creates a new doc with the updated content.
    """

    document = Document(input_filename)
    new_document = Document()
    pbar = tqdm(total=len(document.paragraphs))

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            if len(block.runs) == 0:
                continue
            
            new_paragraph = new_document.add_paragraph()

            new_run = new_paragraph.add_run(block.text)
            new_run.font.name = block.runs[0].font.name  # Preserve font
            new_run.font.size = block.runs[0].font.size  # Preserve font size
            new_run.bold = block.runs[0].bold  # Preserve bold formatting
            new_run.italic = block.runs[0].italic  # Preserve italic formatting

            new_run = new_paragraph.add_run('\n'+translate_text(block.text, target_language))
            new_run.font.name = block.runs[0].font.name  # Preserve font
            new_run.font.color.rgb = RGBColor(0,128,0)
            new_run.font.size = block.runs[0].font.size  # Preserve font size
            new_run.bold = block.runs[0].bold  # Preserve bold formatting
            new_run.italic = block.runs[0].italic  # Preserve italic formatting

        elif isinstance(block, Table):
            new_table = new_document.add_table(rows=len(block.rows)*2, cols=len(block.columns))
            new_table.style = "Table Grid"

            # make a 2d array with len(block.rows) len(block.columns) dimensions and save some values in it
            table_values = [[None for _ in range(len(block.columns))] for _ in range(len(block.rows))]

            for i,row in enumerate(block.rows):
                for j,cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        new_paragraph = cell.paragraphs[0]  # Access existing paragraph
                        for run in paragraph.runs:
                            if table_values[i][j] is None:
                                table_values[i][j] = [run.text + "\n", run.font.name, run.font.size, run.bold, run.italic]
                            table_values[i][j][0] = table_values[i][j][0]+run.text+"\n"

                    table_values[i][j][0]=table_values[i][j][0].strip()

            for i,row in enumerate(new_table.rows):
                for j, cell in enumerate(row.cells):
                    new_i = i/2
                    # convert new_i to int
                    new_i = int(new_i)
                    if i%2==0:
                        cell.text = table_values[new_i][j][0]
                    else:
                        cell.text = '\n'+translate_text(table_values[new_i][j][0], target_language)

        pbar.update(1)
        

    pbar.close()
    new_document.save(output_filename)

# Read the secret key from the .env file
def read_secret_key():
    # Load the environment variables from the .env file
    env_vars = dotenv_values('.env')
    # Access the SECRET_KEY variable
    secret_key = env_vars.get('SECRET_KEY')
    return secret_key

if __name__ == '__main__':
    # Check if the correct number of command line arguments is provided
    if len(sys.argv) < 3:
        print("Usage: python script.py input_file output_file [target_language]")
        sys.exit(1)

    # Extract the input and output file names from command line arguments
    input_file = sys.argv[1]
    output_file = sys.argv[2]

    # Extract the target language if provided,default to 'en' if not specified
    target_language = sys.argv[3] if len(sys.argv) > 3 else 'hindi'

    # Call the function to output the translation
    process_doc(input_file, output_file, target_language)