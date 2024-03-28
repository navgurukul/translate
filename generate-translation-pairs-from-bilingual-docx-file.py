'''
Docstring for Python Script to Process Bilingual Docx File

This Python script processes a docx file containing alternating paragraphs of English and Hindi text. It creates a tab-separated values (TSV) file for further analysis.

Functionality:

Reads the paragraphs from the docx file.
Identifies the language of each paragraph using a character-based detection method.
Creates a TSV file with two columns:
Column 1: English Paragraph (empty if the paragraph is Hindi)
Column 2: Hindi Paragraph (empty if the paragraph is English)
Language Detection:

The script employs a basic language detection technique that relies on character sets. It assumes that paragraphs containing mostly characters from the Devanagari script (used for Hindi) are Hindi, and paragraphs with characters primarily from the Latin alphabet are English. This is a simplified approach and may not be accurate for code-mixed text or paragraphs with special characters.

Note: This method is not foolproof and may require further refinement for handling complex scenarios.
'''

import docx
import os
import sys
import re
import langid
import argparse
from tokenizer import get_sentences_en, get_sentences_hi

def clean_paragraphs(paragraphs):
    # map the above paragraphs text to a new variable paragraphs_text
    paragraphs = [paragraph.text.strip() for paragraph in paragraphs]
    paragraphs = [paragraph.replace('\t','') for paragraph in paragraphs]
    
    # map the paragraphs to remove the first character of the paragraphs beginning with — or -
    paragraphs = [re.sub(r'^—', '', paragraph) for paragraph in paragraphs]
    paragraphs = [re.sub(r'^-', '', paragraph) for paragraph in paragraphs]

    # filter out the empty paragraphs
    paragraphs = list(filter(lambda x: x != '', paragraphs))
    paragraphs = list(filter(lambda x: x != '—', paragraphs))

    return paragraphs

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input_file_en", help="input file english")
    parser.add_argument("input_file_hi", help="input file hindi")
    parser.add_argument("output_file", help="output file")
    args = parser.parse_args()

    input_file_en = args.input_file_en
    input_file_hi = args.input_file_hi
    output_file = args.output_file

    doc_en = docx.Document(input_file_en)
    doc_hi = docx.Document(input_file_hi)

    paragraphs_en = clean_paragraphs(doc_en.paragraphs)
    paragraphs_hi = clean_paragraphs(doc_hi.paragraphs)

    g = open("unmatched.tsv", 'w', encoding='utf-8')

    with open(output_file, 'w', encoding='utf-8') as f:  # Specify encoding as utf-8

        # Initialize a variable i to 0
        i = 0

        # Add an assertion to verify that the length of paragraphs_en is equal to the length of paragraphs_hi and print the length of the paragraphs if the assertion fails
        assert len(paragraphs_en) == len(paragraphs_hi), "Length of paragraphs_en is not equal to length of paragraphs_hi. Length of paragraphs_en is " + str(len(paragraphs_en)) + " and length of paragraphs_hi is " + str(len(paragraphs_hi))
    
        # Loop while i is less than max_len
        while i < len(paragraphs_en):
            # Do the same operations as in the for loop
            en_sentences = get_sentences_en(paragraphs_en[i])
            hi_sentences = get_sentences_hi(paragraphs_hi[i])

            if len(en_sentences) == len(hi_sentences):
                for j in range(len(en_sentences)):
                    f.write(en_sentences[j] + '\t' + hi_sentences[j]+'\n')
            else:
                # write these cases in a unmatched.tsv file
                for j in range(max(len(en_sentences),len(hi_sentences))):
                    if j < min(len(en_sentences), len(hi_sentences)):
                        g.write(en_sentences[j] + '\t' + hi_sentences[j]+'\n')
                    if j >= len(en_sentences):
                        g.write('\t' + hi_sentences[j]+'\n')
                    if j >= len(hi_sentences):
                        g.write(en_sentences[j] + '\t'+'\n')

                g.write('\n')
                # print(en_sentences, hi_sentences,'\n')
            
            i += 1
    
    g.close()

if __name__ == '__main__': 
    main()