from docx import Document

def duplicate_paragraphs(input_path, output_path):
    # Load the input document
    doc = Document(input_path)

    # Create a new document for the output
    output_doc = Document()

    # Iterate over paragraphs in the input document
    for paragraph in doc.paragraphs:
        # Add the original paragraph to the output document
        output_doc.add_paragraph(paragraph.text)

        # Add the duplicated paragraph to the output document
        output_doc.add_paragraph(paragraph.text)

    # Save the output document
    output_doc.save(output_path)

# Provide the paths for the input and output files
input_file = 'input.docx'
output_file = 'output.docx'

# Call the function to duplicate paragraphs
duplicate_paragraphs(input_file, output_file)
