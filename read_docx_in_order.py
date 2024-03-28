#!/usr/bin/env python
# encoding: utf-8

from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

import docx
import random

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)



def process_doc(input_filename, output_filename):
    """
    Reads a docx file in order, replacing characters with spaces randomly,
    and creates a new doc with the updated content.
    """

    document = docx.Document(input_filename)
    new_document = docx.Document()

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            new_paragraph = new_document.add_paragraph()
            for run in block.runs:
                new_run = new_paragraph.add_run(randomize_text(run.text))
                new_run.font.name = run.font.name  # Preserve font
                new_run.font.size = run.font.size  # Preserve font size
                new_run.bold = run.bold  # Preserve bold formatting
                new_run.italic = run.italic  # Preserve italic formatting
        elif isinstance(block, Table):
            new_table = new_document.add_table(rows=len(block.rows), cols=len(block.columns))
            new_table.style = "Table Grid"

            # make a 2d array with len(block.rows) len(block.columns) dimensions and save some values in it
            table_values = [[None for _ in range(len(block.columns))] for _ in range(len(block.rows))]

            for i,row in enumerate(block.rows):
                for j,cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        new_paragraph = cell.paragraphs[0]  # Access existing paragraph
                        for run in paragraph.runs:
                            if table_values[i][j] is None:
                                table_values[i][j] = [run.text + "\n", run.font.name, run.font.size, run.bold, run.italic]
                            table_values[i][j][0] = table_values[i][j][0]+run.text+"\n"

                    table_values[i][j][0]=table_values[i][j][0].strip()

            for i,row in enumerate(new_table.rows):
                for j, cell in enumerate(row.cells):
                    cell.text = table_values[i][j][0]

              
    new_document.save(output_filename)

def randomize_text(text):
    """
    Replaces characters with spaces randomly with a 20% probability.
    """
    new_text = ""
    for char in text:
        if random.random() < 0.2:
            new_text += char
        else:
            new_text += char
    return new_text

# Example usage:
input_filename = "sample.docx"
output_filename = "sample.output.docx"
process_doc(input_filename, output_filename)